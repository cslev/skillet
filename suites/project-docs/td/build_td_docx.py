#!/usr/bin/env python3
"""Build a styled TD .docx from a markdown draft.

Pipeline: pandoc converts the markdown body (+ a Word TOC field) to an
intermediate .docx; python-docx then wraps it with a centered cover page
(logo, title, authors, a "Technology Disclosure (TD)" subtitle, the institute
name, and the date at the bottom), restarts page numbering at the Table of
Contents, justifies the body, and
adds a header (title left / logo right, bottom border) and footer
(page number centered, top border) on every non-cover page.

Two sections are used so the cover has no header/footer and is excluded from
the page count: section 1 = cover, section 2 = TOC + body (numbering restarts
at 1 so the cover never counts and the TOC is page 1).
"""
import argparse
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Emu, RGBColor


def titlecase(s):
    small = {"a", "an", "and", "as", "at", "but", "by", "for", "in", "of",
             "on", "or", "the", "to", "via", "vs", "with"}
    words = s.split()
    out = []
    for i, w in enumerate(words):
        lw = w.lower()
        if 0 < i < len(words) - 1 and lw in small:
            out.append(lw)
        elif w.isupper() and len(w) > 1:   # keep acronyms (TD, GPU) as-is
            out.append(w)
        else:
            out.append(w[:1].upper() + w[1:])
    return " ".join(out)


def strip_front_matter(md_text):
    """Drop everything before the first 'Abstract' heading. The title and
    author block live on the cover page, not in the body/TOC."""
    lines = md_text.splitlines()
    for i, line in enumerate(lines):
        m = re.match(r"^#{1,3}\s+(.*\S)\s*$", line)
        if m and m.group(1).strip().lower() == "abstract":
            return "\n".join(lines[i:])
    return md_text  # no Abstract heading found; use whole file


def field_run(paragraph, instr):
    """Append a Word field (e.g. PAGE) that updates on open."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, end):
        run._r.append(el)


def set_border(paragraph, edge):
    """Single black border on one edge of a paragraph (header/footer rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "6")        # 0.75 pt
    el.set(qn("w:space"), "4")
    el.set(qn("w:color"), "000000")
    pBdr.append(el)


def copy_page_geometry(src_sectPr, dst_sectPr):
    for tag in ("w:pgSz", "w:pgMar"):
        src = src_sectPr.find(qn(tag))
        if src is not None:
            import copy
            dst_sectPr.append(copy.deepcopy(src))


def build(md_path, out_path, title, authors, date, logo, institute):
    # Resolve everything to absolute paths up front. pandoc resolves relative
    # image links against its working directory (not the input file's), so we
    # run it from the draft's own directory and figure paths written relative
    # to the draft (e.g. ../assets/fig.png) resolve correctly regardless of
    # where this script is invoked from. The logo is embedded by python-docx
    # (this process's CWD), so it must be absolute too.
    md_path = Path(md_path).resolve()
    md_dir = md_path.parent
    out_path = str(Path(out_path).resolve())
    if logo:
        logo = str(Path(logo).resolve())
    md_text = md_path.read_text(encoding="utf-8")
    body_md = strip_front_matter(md_text)

    with tempfile.TemporaryDirectory() as tmp:
        body_md_path = Path(tmp) / "body.md"
        body_md_path.write_text(body_md, encoding="utf-8")
        body_docx = Path(tmp) / "body.docx"
        proc = subprocess.run(
            ["pandoc", str(body_md_path), "-o", str(body_docx),
             "--toc", "--toc-depth=3"],
            cwd=str(md_dir),
            capture_output=True,
            text=True,
        )
        if proc.returncode != 0:
            sys.stderr.write(proc.stderr)
            raise SystemExit(f"pandoc failed (exit {proc.returncode})")
        # pandoc drops unresolved images with only a warning and exit 0 — make
        # that loud so missing figures are never shipped silently.
        missing = [ln for ln in proc.stderr.splitlines()
                   if "Could not fetch resource" in ln]
        if missing:
            sys.stderr.write(
                "\nWARNING: pandoc could not fetch these images; the .docx will "
                "be missing figures. Check the paths (they must be relative to "
                f"the draft at {md_dir}):\n")
            for ln in missing:
                sys.stderr.write("  " + ln + "\n")
        doc = Document(str(body_docx))

    # --- ensure explicit page geometry (pandoc leaves it inherited) --------
    sec0 = doc.sections[0]
    if sec0.page_width is None:
        sec0.page_width = Inches(8.5)
    if sec0.page_height is None:
        sec0.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        if getattr(sec0, attr) is None:
            setattr(sec0, attr, Inches(1))

    # --- justify body text -------------------------------------------------
    for style_name in ("Normal", "Body Text", "First Paragraph"):
        try:
            doc.styles[style_name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass

    body = doc.element.body
    first = body[0]  # the pandoc TOC sdt

    # --- page breaks: TOC, Abstract, main content, and References each
    #     start on a fresh page ---------------------------------------------
    def page_break_before(paragraph):
        paragraph.insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)

    h1s = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    for idx, p in enumerate(h1s):
        heading = p.text.strip().lower()
        if idx == 0:                    # Abstract -> TOC keeps its own page
            page_break_before(p)
        elif idx == 1:                  # main content starts after the Abstract
            page_break_before(p)
        elif heading == "references":   # References on its own page
            page_break_before(p)

    # --- References entries are left-aligned; the rest stays justified -----
    in_refs = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading 1"):
            in_refs = p.text.strip().lower() == "references"
            continue
        if in_refs:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- cover page (inserted before the TOC) ------------------------------
    def add_before(make):
        """Create an element via the high-level API (appended at the end of
        the body) then move it to just before `first`, preserving call order."""
        make()
        el = body[-2] if body[-1].tag == qn("w:sectPr") else body[-1]
        first.addprevious(el)
        return el

    if logo:
        def _logo():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(72)
            p.add_run().add_picture(logo, width=Inches(2.4))
        add_before(_logo)

    def _title():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(48 if not logo else 24)
        p.paragraph_format.space_after = Pt(24)
        r = p.add_run(titlecase(title))
        r.bold = True
        r.font.size = Pt(26)
    add_before(_title)

    def _authors():
        n = len(authors)
        t = doc.add_table(rows=n, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for row, name in zip(t.rows, authors):
            cell = row.cells[0]
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(name)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    add_before(_authors)

    def _subtitle():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(28)
        r = p.add_run("Technology Disclosure (TD)")
        r.italic = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    add_before(_subtitle)

    if institute:
        def _institute():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(institute)
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        add_before(_institute)

    # The date is pinned to the bottom of the cover via the cover section's
    # footer (set up below). Putting it in the body and pushing it down with
    # spacing is unreliable: it overflows onto a second page when the title
    # wraps or there are many authors.

    # --- section break: carrier paragraph ends the cover section ----------
    body_sectPr = body[-1]
    carrier = doc.add_paragraph()
    carrier_p = carrier._p
    first.addprevious(carrier_p)
    pPr = carrier_p.get_or_add_pPr()
    cover_sectPr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    cover_sectPr.append(sect_type)
    copy_page_geometry(body_sectPr, cover_sectPr)
    pPr.append(cover_sectPr)

    # --- cover section footer: the date, centered, no page number ---------
    cover_section = doc.sections[0]
    cfp = cover_section.footer.paragraphs[0]
    cfp.text = ""
    cfp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cfp.add_run(date)
    crun.font.size = Pt(12)

    # --- body section: restart numbering, header + footer ------------------
    body_section = doc.sections[1]

    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    body_sectPr.append(pgNumType)

    usable = body_section.page_width - body_section.left_margin - body_section.right_margin

    header = body_section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.tab_stops.add_tab_stop(Emu(usable), WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run(titlecase(title))
    if logo:
        hp.add_run("\t")
        hp.add_run().add_picture(logo, height=Inches(0.3))
    set_border(hp, "bottom")

    footer = body_section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field_run(fp, "PAGE")
    set_border(fp, "top")

    doc.save(out_path)
    print(f"Wrote {out_path}  (sections={len(doc.sections)})")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--title", required=True)
    ap.add_argument("--authors", required=True,
                    help="Author names separated by ';'")
    ap.add_argument("--date", required=True)
    ap.add_argument("--institute", required=True,
                    help="Institute / organization name (confirmed by the user)")
    ap.add_argument("--logo", default="", help="Path to a logo image, or empty")
    a = ap.parse_args()
    authors = [s.strip() for s in a.authors.split(";") if s.strip()]
    build(a.md, a.output, a.title, authors, a.date, a.logo or None, a.institute)


if __name__ == "__main__":
    sys.exit(main())
